#!/usr/bin/env python3
"""
Markdown 백서 → DOCX 변환 스크립트 (개별 구동).

사용법:
    python scripts/md_to_docx.py output.md
    python scripts/md_to_docx.py output.md -o report.docx
    python scripts/md_to_docx.py output.md output_en.md -o combined.docx

복수 파일 전달 시 순서대로 하나의 DOCX에 병합 (페이지 구분).

의존성:
    pip install python-docx
"""
from __future__ import annotations
import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[ERROR] python-docx 미설치. 설치 후 재실행:")
    print("  pip install python-docx")
    sys.exit(1)


# ──────────────────────────────────────────────────────────────
# Style constants
# ──────────────────────────────────────────────────────────────
FONT_NAME = "맑은 고딕"
FONT_NAME_EN = "Calibri"
FONT_SIZE_BODY = Pt(10.5)
FONT_SIZE_H1 = Pt(24)
FONT_SIZE_H2 = Pt(15)
FONT_SIZE_H3 = Pt(12)
LINE_SPACING = Pt(18)
PAGE_MARGIN = Cm(2.54)

# Brand colors (business report palette)
COLOR_PRIMARY = RGBColor(0x1F, 0x3A, 0x5F)   # deep navy — headings
COLOR_ACCENT = RGBColor(0x2E, 0x6D, 0xB4)    # blue — accents
COLOR_BODY = RGBColor(0x22, 0x22, 0x22)      # near-black body
COLOR_MUTED = RGBColor(0x70, 0x70, 0x70)     # gray — meta
COLOR_RULE = "1F3A5F"
COLOR_CALLOUT_BG = "EEF3FA"                   # light blue callout fill


# ──────────────────────────────────────────────────────────────
# Markdown parsing
# ──────────────────────────────────────────────────────────────
_RE_H1 = re.compile(r'^# (.+)$')
_RE_H2 = re.compile(r'^## (.+)$')
_RE_H3 = re.compile(r'^### (.+)$')
_RE_BULLET = re.compile(r'^- (.+)$')
_RE_BLOCKQUOTE = re.compile(r'^> (.+)$')
_RE_BOLD = re.compile(r'\*\*(.+?)\*\*')
_RE_INLINE_CODE = re.compile(r'`([^`]+)`')
_RE_HR = re.compile(r'^---+$')


def _add_styled_runs(paragraph, text: str):
    """Parse inline markdown (bold, inline code) into styled runs."""
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        bold_match = _RE_BOLD.fullmatch(part)
        code_match = _RE_INLINE_CODE.fullmatch(part)
        if bold_match:
            run = paragraph.add_run(bold_match.group(1))
            run.bold = True
        elif code_match:
            run = paragraph.add_run(code_match.group(1))
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        else:
            paragraph.add_run(part)


def _set_cn_font(run, name: str = FONT_NAME):
    """Ensure East-Asian font is applied (python-docx omits w:eastAsia by default)."""
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), name)


def _shade_paragraph(paragraph, fill_hex: str):
    """Apply background shading to a paragraph."""
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    ppr.append(shd)


def _bottom_border(paragraph, color_hex: str, size: int = 12):
    """Add a bottom border (rule) under a paragraph."""
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pbdr.append(bottom)
    ppr.append(pbdr)


def _setup_doc_styles(doc: Document):
    """Configure document-level styles and page layout."""
    # Page margins
    for section in doc.sections:
        section.top_margin = PAGE_MARGIN
        section.bottom_margin = PAGE_MARGIN
        section.left_margin = PAGE_MARGIN
        section.right_margin = PAGE_MARGIN

    # Default font
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE_BODY
    style.font.color.rgb = COLOR_BODY
    style.paragraph_format.line_spacing = LINE_SPACING
    style.paragraph_format.space_after = Pt(6)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), FONT_NAME)


def _apply_heading_font(paragraph, size: Pt):
    """Apply heading font styling."""
    for run in paragraph.runs:
        _set_cn_font(run)
        run.font.size = size
        run.bold = True
        run.font.color.rgb = COLOR_ACCENT


def _add_cover_title(doc: Document, title: str):
    """Render a polished cover-style document title (H1)."""
    # Top label
    label = doc.add_paragraph()
    label.paragraph_format.space_after = Pt(2)
    lr = label.add_run("BUSINESS WHITEPAPER")
    lr.font.size = Pt(9)
    lr.bold = True
    lr.font.color.rgb = COLOR_ACCENT
    _set_cn_font(lr)

    # Title
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title)
    _set_cn_font(run)
    run.font.size = FONT_SIZE_H1
    run.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    _bottom_border(p, COLOR_RULE, size=18)

    # Spacer after title block
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)


def md_to_docx(md_text: str, doc: Document, title_label: str = ""):
    """Convert markdown text to DOCX content in the given document."""

    if title_label:
        p = doc.add_paragraph()
        run = p.add_run(title_label)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    in_blockquote = False
    blockquote_lines: list = []

    def _flush_blockquote():
        nonlocal in_blockquote, blockquote_lines
        if blockquote_lines:
            text = " ".join(blockquote_lines)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            _add_styled_runs(p, text)
            # Light gray left border effect via indentation
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                run.font.size = Pt(9.5)
        blockquote_lines = []
        in_blockquote = False

    for line in md_text.split("\n"):
        stripped = line.strip()

        # Skip empty lines (flush blockquote if active)
        if not stripped:
            if in_blockquote:
                _flush_blockquote()
            continue

        # Horizontal rule → thin separator rule (not a page break)
        if _RE_HR.match(stripped):
            if in_blockquote:
                _flush_blockquote()
            rule = doc.add_paragraph()
            rule.paragraph_format.space_before = Pt(6)
            rule.paragraph_format.space_after = Pt(6)
            _bottom_border(rule, "C9D4E2", size=6)
            continue

        # Blockquote
        bq = _RE_BLOCKQUOTE.match(stripped)
        if bq:
            in_blockquote = True
            blockquote_lines.append(bq.group(1))
            continue
        elif in_blockquote:
            _flush_blockquote()

        # H1 (document title) → cover-style title
        h1 = _RE_H1.match(stripped)
        if h1:
            _add_cover_title(doc, h1.group(1))
            continue

        # H2
        h2 = _RE_H2.match(stripped)
        if h2:
            p = doc.add_heading("", level=2)
            run = p.add_run(h2.group(1))
            _set_cn_font(run)
            run.font.size = FONT_SIZE_H2
            run.bold = True
            run.font.color.rgb = COLOR_PRIMARY
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            _bottom_border(p, COLOR_RULE, size=8)
            continue

        # H3
        h3 = _RE_H3.match(stripped)
        if h3:
            p = doc.add_heading(h3.group(1), level=3)
            _apply_heading_font(p, FONT_SIZE_H3)
            continue

        # Bullet list
        bullet = _RE_BULLET.match(stripped)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            _add_styled_runs(p, bullet.group(1))
            for run in p.runs:
                _set_cn_font(run)
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_styled_runs(p, stripped)
        for run in p.runs:
            _set_cn_font(run)

    # Flush remaining blockquote
    if in_blockquote:
        _flush_blockquote()


def _add_page_numbers(doc: Document):
    """Add centered 'page / total' numbering to the footer."""
    def _field(instr: str):
        fld = OxmlElement('w:fldSimple')
        fld.set(qn('w:instr'), instr)
        return fld

    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.text = ""
        run = p.add_run()
        run.font.size = Pt(8)
        run.font.color.rgb = COLOR_MUTED
        p._p.append(_field('PAGE'))
        sep = p.add_run(" / ")
        sep.font.size = Pt(8)
        sep.font.color.rgb = COLOR_MUTED
        p._p.append(_field('NUMPAGES'))


def build_whitepaper_docx(md_text: str, out_path):
    """최종 마크다운 백서 → 세련된 비즈니스 보고서 DOCX로 변환.

    제목(H1) + 본문(H2 섹션) + 시사점 구조를 그대로 반영.
    main.py 에서 파이프라인 종료 시 자동 호출된다.
    """
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _setup_doc_styles(doc)
    md_to_docx(md_text, doc, title_label="")
    _add_page_numbers(doc)
    doc.save(str(out_path))
    return out_path



# ──────────────────────────────────────────────────────────────
# CLI
# ──────────────────────────────────────────────────────────────
def parse_args():
    p = argparse.ArgumentParser(
        description="Markdown 백서 → DOCX 변환",
        epilog="복수 파일 전달 시 순서대로 병합 (페이지 구분)",
    )
    p.add_argument("inputs", nargs="+", help="입력 마크다운 파일 (1개 이상)")
    p.add_argument("-o", "--output", default=None,
                   help="출력 DOCX 경로 (기본: 첫 입력 파일명.docx)")
    return p.parse_args()


def main():
    args = parse_args()

    # Validate inputs
    input_paths = []
    for inp in args.inputs:
        p = Path(inp)
        if not p.exists():
            print(f"[ERROR] 파일 없음: {p}")
            sys.exit(1)
        input_paths.append(p)

    # Output path
    if args.output:
        out_path = Path(args.output)
    else:
        out_path = input_paths[0].with_suffix(".docx")

    # Build document
    doc = Document()
    _setup_doc_styles(doc)

    for i, inp in enumerate(input_paths):
        if i > 0:
            doc.add_page_break()

        md_text = inp.read_text(encoding="utf-8")
        label = inp.name if len(input_paths) > 1 else ""
        md_to_docx(md_text, doc, title_label=label)
        print(f"  [{i+1}/{len(input_paths)}] {inp.name} ({len(md_text):,} chars)")

    doc.save(str(out_path))
    print(f"\n✅ 저장: {out_path.resolve()}")


if __name__ == "__main__":
    main()
